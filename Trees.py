import networkx as nx
import pylab
limits=pylab.axis('off')
import matplotlib.pyplot as plt
import time
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE

'''
=============================================
    PART I: useful background functions
=============================================
'''

def roman(n): #takes a number and gives back the Roman numeral version
    if n==0:
        return"Zero"
    anums = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    rnums ="M CM D CD C XC L XL X IX V IV I".split()
    ret = []
    for a,r in zip(anums, rnums):
        x,n = divmod(n,a)
        ret.append(r*x)
    return ''.join(ret)
        

def scale(n): #takes a number and gives back the largest number with a factorial less than the input - so, sort of the floor of the inverse factorial
    s=0
    m=n
    while m>s:
        s=s+1
        m=m//s
    return s

def factorial(n): # takes a number and gives back its factorial
    if n<=1:
        return 1
    else:
        return n*factorial(n-1)

def fbase(tree,size=0): # takes a number and a size, and gives back a factoradic representation of the number, padded with zeroes if needed.
    m=tree
    s=scale(tree)
    if size<s:
        size=s
    #s=s-1
    f=[]
    fs=factorial(s)
    for t in range(s,0,-1):
        q=int(m//fs)
        f.append(q)
        d=int(fs*q)
        m=int(m-d)
        fs=int(fs/t)
    f.append(0)
    return f

def dbase(fb):# takes a factoradic representation of a number, and gives back the decimal form of the number.
    size=len(fb)
    db=0
    s=0
    s+=size-1
    f=factorial(s)
    while s != 0:
        a=fb.pop(0)
        db+=a*f
        f=f//s
        s=s-1
    return db
'''
=============================================
    PART I: basic tree functions
=============================================
'''

def graceful(tree,size=0): # takes an index number and a size of tree and gives back a set of links for a graceful graph
    s=scale(tree)
    if size>=s:
        s=size
    f=fbase(tree,s)
    w=len(f)
    g=[]
    for i in range(0,w):
        g.append([f[i],f[i]+i+1])
    return g

def adjacency(tree,size=0): # takes an index number and a size of tree and gives back a printout of the adjacency matrix
    s=scale(tree)
    if size>s:
        s=size
    s=s-1
    g=graceful(tree,size)
    adj=[]
    for i in range(0,s+3):
        adj.append("|")
        for j in range(0,i):
            if [j,i] in g:
                adj[i] +="@|"
            else:
                adj[i] +=" |"
        print(adj[i])
        
def valence(tree,size=0): # takes an index number and a size of tree and gives back the vector of node valences (or orders)
    s=scale(tree)
    if s>size:
        size=s
    graph=expansion(graceful(tree,size))
    v=[]
    for node in graph:
        v.append(0)
    for node in graph:
        v[len(node)]+=1
    
    '''f=fbase(tree,s)
    w=len(f)
    c=[]
    v=[]
    for i in range(0,w):
        f.append(f[i]+i+1)
        print(f)
    for i in range(0,w+1):
        c.insert(0,f.count(i))
        print(c)
    for i in range(0,w+1):
        v.append(c.count(i))
        print(v)'''
    return v
   
def draw(pg,tree,size=0): # takes a tree species (number from 0 to the size of the vlist),an index number and a size of tree, and saves a picture of the tree.
    s=scale(tree)
    if s<size:
        s=size
    graph=graceful(tree,s)
    path=trunk(tree,s)
    G = nx.Graph()
    for q in graph:
        G.add_edge(q[0],q[1])
    red_edges=[]
    black_edges=[]
    for edge in graph:
        if edge[0] in path and edge[1] in path:
            red_edges.append(edge)
        else:
            black_edges.append(edge)            
    edge_colours = ['black' if not edge in red_edges else 'red' for edge in G.edges()]
    val_map={0:'palegreen'}
    values=[val_map.get(node, 'skyblue') for node in G.nodes()]
    pos = nx.fruchterman_reingold_layout(G)
    nx.draw_networkx_nodes(G, pos, cmap=plt.get_cmap('jet'), node_color=values, node_size = 2000)
    nx.draw_networkx_labels(G, pos)
    nx.draw_networkx_edges(G, pos, edgelist=red_edges, edge_color='r')
    nx.draw_networkx_edges(G, pos, edgelist=black_edges)
    plt.savefig("C:/Users/Alec/OneDrive/Python/Journal/tree-"+str(s)+"-"+str(pg)+"-"+str(tree)+".png")
    plt.close()

def display(tree,size=0): # takes a tree species (number from 0 to the size of the vlist),an index number and a size of tree, and displays a picture of the tree.
    s=scale(tree)
    if s<size:
        s=size
    graph=graceful(tree,s)
    path=trunk(tree,s)
    G = nx.Graph()
    for q in graph:
        G.add_edge(q[0],q[1])
    red_edges=[]
    black_edges=[]
    for edge in graph:
        if edge[0] in path and edge[1] in path:
            red_edges.append(edge)
        else:
            black_edges.append(edge)            
    edge_colours = ['black' if not edge in red_edges else 'red' for edge in G.edges()]
    val_map={0:'palegreen'}
    values=[val_map.get(node, 'skyblue') for node in G.nodes()]
    pos = nx.fruchterman_reingold_layout(G)
    nx.draw_networkx_nodes(G, pos, cmap=plt.get_cmap('jet'), node_color=values, node_size = 2000)
    nx.draw_networkx_labels(G, pos)
    nx.draw_networkx_edges(G, pos, edgelist=red_edges, edge_color='r')
    nx.draw_networkx_edges(G, pos, edgelist=black_edges)
    plt.show()



def nxt(m,graph):# takes a node label and a graph, and returns the link with the first occurrence of that label
    w=len(graph)
    for i in range(0,w):
        if graph[i][0]==m or graph[i][1]==m:
            return i
    return -1

def jnt(m,graph): # takes a node label and a graph, and returns other node in the link with the first occurrence of that label
    w=len(graph)
    for i in range(0,w):
        if graph[i][0]==m:
            return graph[i][1]
        elif graph[i][1]==m:
            return graph[i][0]
    return -1

def connect(tree,size=0): # takes an index number and a size of tree and says whether the graph is a tree
    graph=graceful(tree,size)
    links=[0]
    w=len(graph)
    t=w*(w+1)/2
    while links !=[]:
        n=jnt(links[0],graph)
        i=nxt(links[0],graph)
        if n!= -1:
            if not n in links:
                links.append(n)
            graph.pop(i)
        else:
            t=t-links[0]
            links.pop(0)
    boo=(len(graph)==0 and t==0)
    return boo

def card(tree,size=0): # takes an index number and a size of tree and prints out basic information - its links, adjacency matrix etc, and displays the graph
    s=scale(tree)
    if size>s:
        s=size
    print('===================================')
    print(tree,' - ',graceful(tree,s))
    adjacency(tree,s)
    if connect(tree,s):
        print('Connected tree')
    else:
        print('Non-connected graph')
    print(valence(tree,s))
    print('trunk = ',trunk(tree,s))
    display(tree,size)
        
def cards(k): #takes the size of the trees and prints out cards for all graceful trees of that size.
    m=0
    n=factorial(k+1)
    for i in range(0,n,2):
        if connect(i,k):
            m+=1
            card(i,k)
    return m

def matrix(document,tree,size=0):#inserts the adjacency matrix for a given tree into the given document, including a space to insert a graphic of the tree.
    s=scale(tree)
    if s>size:
        size=s
    g=graceful(tree,size)
    table = document.add_table(rows=size+2, cols=size+3)
    table.autofit = True
    for i in range(0,size+1):
        for cell in table.columns[i].cells:
            cell.width = Inches(0.1)
    for i in range(0,size+1):
        entry=table.cell(i,i)
        para=entry.add_paragraph(str(i))
        para.style=document.styles['No Spacing']
        for j in range(0,i):
            if [j,i] in g:
                shading_elm = parse_xml(r'<w:shd {} w:fill="111111"/>'.format(nsdecls('w')))
            else:
                shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.cell(i,j)._tc.get_or_add_tcPr().append(shading_elm)
    
    left=table.cell(size+1,0)
    right=table.cell(size+1,size+1)
    b=left.merge(right)
    top=table.cell(0,size+2)
    bottom=table.cell(size+1,size+2)
    a=top.merge(bottom)
    return table

def contents(rank): #Produces an MS Word document listing all graceful trees with k links, organised by class-number r.
                    #The 'class-number' is the position in 'catalogue' (qv) of a given valence-vector.
    size=rank
    v=vlist(size)
    cat=catalogue(size)
    w=len(cat)
    document=Document()
    document.add_heading('Graceful Trees of Rank ' + roman(rank),level=0).bold=True
    document.add_heading("Contents",level=1).bold=True
    items = 0
    for line in range(0,w):
        kind=cat[line]
        items += len(cat[line])
        document.add_paragraph("Class " + str(line) +": " + str(valence(kind[0],rank)))
        if len(cat[line])==1:
            word=" tree, "
        else:
            word=" trees, "
        document.add_paragraph(str(len(cat[line])) + word + str(cat[line]))
        document.add_paragraph("")
    document.add_paragraph(str(items) +" trees in total.").bold=True
    document.add_paragraph("(Remember that each featured tree has a corresponding odd-numbered partner, called its 'converse', where every label x is replaced with " + str(rank+1) +"-x, and with a tree-index number of " + str(factorial(rank)-1) +" - x)").italic=True
    document.save("C:/Users/Alec/OneDrive/Python/Journal/Contents of Rank " + roman(rank) +".docx")


def journal(rank,sort=-1):  # Produces an MS Word document listing each graceful tree with given rank, of class-number 'sort'; along with its list of links,
                            # its adjacency vector and a picture of the tree.
                            # The class-number is the position in 'catalogue' (qv) of a given valence-vector. 
    size=rank
    if sort==-1:
        sub=""
    else:
        sub=" - " + str.lower(roman(sort))
    v=vlist(size)
    cat=catalogue(size)
    w=len(cat)
    if w<sort:
        return "Class Number out of range (max =" + str(w) + ")"
    if sort==-1:
        a=0
        b=w
    else:
        a=sort
        b=sort+1
    document=Document()
    document.add_heading('Graceful Trees of Rank ' + roman(rank),level=0).bold=True
    if sort!=-1:
        document.add_heading('Class ' + roman(sort), level=1)
    document.add_page_break()
    document.add_heading("Contents",level=2).bold=True
    items = 0
    for line in range(a,b):
        kind=cat[line]
        items += len(cat[line])
        document.add_paragraph("Class" + str(line) +": " + str(valence(kind[0],rank)))
        if len(cat[line])==1:
            word==" tree, "
        else:
            word=" trees, "
        document.add_paragraph(str(len(cat[line])) + word + str(cat[line]))
        document.add_paragraph("")
    document.add_paragraph(str(items) +" trees in total.")
    document.add_paragraph("(Remember that each featured tree has a corresponding odd-numbered partner, called its 'converse', where every label x is replaced with " + str(rank+1) +"-x, and with a tree-index number of " + str(factorial(rank)-1) +" - x)")
    m=0
    n=factorial(rank)
    for pos in range(a,b):
        kind=cat[pos]
        document.add_page_break()
        document.add_heading("Class" + str(pos) +":  " + str(valence(kind[0],rank)),level=2)
            
        for tree in kind:
            if connect(tree,size):
                document.add_heading("Tree " + str(tree),level=3)
                document.add_paragraph('Links = ' + str(graceful(tree,size)))
                s=scale(tree)
                if s>size:
                    size=s
                g=graceful(tree,size)
                t=matrix(document,tree,size)
                para=t.cell(0,size+2).add_paragraph()
                draw(pos,tree,size)
                para.add_run().add_picture("C:/Users/Alec/OneDrive/Python/Journal/tree-"+str(rank)+"-"+str(pos)+"-"+str(tree)+".png",width=Inches(0.5*size))
                #document.add_paragraph('trunk    =     ' + str(trunk(tree,size)))
                #document.add_paragraph('Branches = ' + str(twig(tree,size)))
                t.cell(s+2,0).add_paragraph('Trunk: ' + str(trunk(tree,size)))
                t.cell(s+2,0).add_paragraph('Twigs: ' + str(twig(tree,size)))
                
        m+=1
    document.save("C:/Users/Alec/OneDrive/Python/Journal/Rank " + roman(rank) + sub +".docx")

def twig(tree,size=0):
    s=scale(tree)
    if size>s:
        s=size
    path=trunk(tree,s)
    twigs=[]
    ex=expansion(graceful(tree,s))
    for node in path:
        n=len(ex[node])
        if n==1:
            twigs.append(0)
        else:
            twigs.append(n-2)
    rt=reverse(twigs)
    if weight(twigs)<weight(rt):
        twigs=rt
    return twigs

def ok(k):#returns the number of non-trees and trees, among all graceful graphs of size k
    ok=[0,0]
    n=factorial(k)
    for i in range(0,n,2):
        if connect(i,k-1):
            ok[1]+=1
        else:
            ok[0]+=1
    return ok

def catalogue(k):#searches all graceful trees of size k to list all their valence vectors.
    n=factorial(k)
    cat=[]
    v=vlist(k)
    for u in v:
        cat.append([])
    for i in range(0,n):
        if connect(i,k-1):
            #print(i,valence(i,k-1),connect(i,k-1))
            w=valence(i,k-1)
            if w in v:
                p=v.index(w)
                cat[p].append(i)
    return cat

def tlist(k):
    n=factorial(k)
    tlist=[]
    for tree in range(0,n):
        if connect(tree,k):
            entry=(valence(tree,k),twig(tree,k))
            if entry not in tlist:
                tlist.append(entry)
    return tlist


def escalate(v,j):#takes a valence vector v and adds a new node linked to a node of order j. If v(j)=0, it returns the input vector unchanged.
    if j not in range(0,len(v)+1):
        return v
    elif j == len(v)-1:
        v.append(0)
    if v[j] == 0:
        return v
    v[j+1]=v[j+1]+1
    v[j]=v[j]-1
    v[1]=v[1]+1
    return v

def vlist(k):#generates all possible valence vectors of size k, by recursively escalating vectors of size k-1.
    if k==0:
        return [[1]]
    ww=vlist(k-1)
    w=ww[:]
    v=[]
    for u in w:
        for i in range(0,len(u)):
            t=u[:]
            ee=escalate(t,i)
            e=ee[:]
            if e != u:
                while len(e)<=k:
                    e.append(0)
                if e not in v:
                    v.append(e)
    for a in v:
        aa=0
        for b in a:
            aa += b
        if aa<=k:
            v.remove(a)
    return v

def kind(valen):#takes a valence vector and returns a list of index numbers of trees that have the same valence vector, taking the sum of the vector components as the size of the trees.
    s=0
    k=[]
    for i in valen:
        s=s+i
    s=s-2
    while len(valen)<s+1:
        valen.append(0)
    for i in range(0,factorial(1+s)):
        g=graceful(i,s)
        v=valence(i,s)
        if v==valen:            
            k.append(i)
    return k

def expansion(graph): # lists all the nodes that link to the nth node, where n is the place in this list.
    ex=[[]]
    for links in graph:
        ex.append([])
    for links in graph:
        ex[links[0]].append(links[1])
        ex[links[1]].append(links[0])
    return ex
        

def trunk(tree,size=0): # creates a list of all the nodes on the longest path or"trunk" of the tree
    s=scale(tree)
    j=scale(tree)
    if s>size:
        size=s
    g=expansion(graceful(tree,size))
    path=[]
    for nodes in g:
        path.append([-1,-1,-1])
        if len(nodes)==1:
            j=g.index(nodes)
    for run in range(0,3):
        mark(j,run,path,g,0)
        mx=0
        for nodes in path:
            if nodes[run]>mx:
                mx=nodes[run]
                j=path.index(nodes)
    p=[]
    for j in range(0,mx+1):
        p.append(0)
    for nodes in path:
        if nodes[1]+nodes[2]==mx:
            p[nodes[2]]=path.index(nodes)
    return p

def mark(j,n,path,graph,count):# recursively marks all the nodes in a tree, to calculate the longest path or"trunk"
    path[j][n]=count
    count+=1
    for k in graph[j]:
        if path[k][n]==-1:
            mark(k,n,path,graph,count)

def reverse(list0):
    k=[]
    for item in list0:
        k.insert(0,item)
    return k
def weight(list0):
    w=0
    for item in list0:
        w=w+item
    return w

def plant(tree,size=0):
    s=scale(tree)
    if s>size:
        size=s
    g=graceful(tree,size)
    e=expansion(g)
    t=trunk(tree,size)
    b=twig(tree,size)
    loc=[]
    for node in t:
        n=t.index(node)
        loc.append([node,[0,size/2-n]])
                                             
        

    
def readme():
    print("")
    print("=============================================================")
    print("TREES: routines to help visualise the Graceful Trees problem.")
    print("=============================================================")
    print("")
    print("All the graphs generated by these functions are graceful; though not all are")
    print("trees. You can use connect(n,<p>) to weed out non-trees. In functions that")
    print("generate graceful trees, only ones with an even index are considered, since")
    print("every odd-indexed tree is the converse or 'flip' of an even-indexed tree. To see")
    print("examples, use card(n,<p>) to view such pairs - the index of the flip of a tree")
    print("is (k+1)!-(n+1), where k is size of the tree. For instance, the flip of (34,4)")
    print("is 5!-35, i.e. (85,4)")
    print("")
    print("Where the parameters of a function are given as (n,<p>), n is the index number")
    print("of the graph, and p is the number of links the graph has. If p is omitted, the")
    print("graph will be the smallest possible graph with that index number.")
    print("")
    print("Some minor functions are ommitted from this list, if their main use is inside")
    print("another function.")
    print("")
    print("")
    print("=========================")
    print("FUNCTIONS")
    print("=========================")
    print("adjacency(n,<p>): prints out the adjacency matrix for that graph.")
    print("")
    print("graceful(n,<p>):  returns a list of links for the corresponding graceful graph.")
    print("")
    print("valence(n,<p>:    returns the valence vector for the graph - item k is the")
    print("                  number of nodes with k links.")
    print("")
    print("draw(n,<p>):      generates a visual representation of the graph, including")
    print("                  location of the zero, and the trunk. Saves it to a file.")
    print("")
    print("display(n,<p>):   generates a visual representation of the graph, including")
    print("                  location of the zero, and the trunk. Displays it, but")
    print("                  doesn't save it.")
    print("")
    print("connect(n,<p>):   returns true if the graph is a tree, or false if the graph")
    print("                  is not connected.")
    print("")
    print("expansion(graph): lists all the nodes that link to the nth node, where n is")
    print("                  the place in this list.")
    print("                 ")
    print("trunk(n,<p>):     creates a list of all the nodes on the longest path or 'trunk'")
    print("                  of the tree.")
    print("                 ")
    print("mark(j,n,p,g,c):  recursively marks all the nodes in a tree, to calculate the")
    print("                  longest path or 'trunk'.")
    print("                 ")
    print("                 ")
    print("=========================")
    print("REPORTS")
    print("=========================")
    print("card(n,<p>):      prints the link list, matrix and valence vector, as well as")
    print("                  drawing the graph.")
    print("                 ")
    print("cards(k):         prints all the cards for graceful trees of size k, in order of")
    print("                  index-number. Advance to the next one by closing the graphics")
    print("                  window.")
    print("")
    print("ok(k):            returns the number of non-trees and trees, among all graceful")
    print("                  graphs of size k.")
    print("                 ")
    print("catalogue(k):     searches all graceful trees of size k to list all their")
    print("                  valence vectors.")
    print("")
    print("vlist(k):         generates all possible valence vectors of size k, by")
    print("                  recursively escalating vectors of size k-1.")
    print("")
    print("kind(v-vector):   takes a valence vector and returns a list of index numbers of")
    print("                  trees that have the same valence vector, taking the sum of the")
    print("                  vector components as the size of the tree.")
    print("                 ")
    print("journal(k,<r>):   produces an MS Word document listing all graceful trees with k")
    print("                  links, of class-number r. The class-number is the position in")
    print("                  'catalogue' (qv) of a given valence-vector. If r is left out,")
    print("                  the document lists all graceful trees with k links, organised")
    print("                  by class-number.")
    print("                 ")
    print("                 ")
    print("=========================")
    print("OPERATORS")
    print("=========================")
    print("escalate(v,j):    takes a valence vector v and adds a new node linked to a node")
    print("                  of order j. If v(j)=0, it returns the input vector unchanged.")
    print("")
